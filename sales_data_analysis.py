#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
제품 판매 데이터 분석 및 보고서 생성 스크립트
Sales Data Analysis and Report Generation Script
"""

import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import warnings
from typing import Optional, Dict, Any
import smtplib
import os
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
warnings.filterwarnings('ignore')

class SalesDataAnalyzer:
    def __init__(self, csv_file_path: str):
        """
        판매 데이터 분석기 초기화
        """
        self.csv_file_path = csv_file_path
        self.df: Optional[pd.DataFrame] = None
        self.cleaned_df: Optional[pd.DataFrame] = None
        self.analysis_results: Dict[str, Any] = {}
        
    def load_and_clean_data(self):
        """
        데이터 로드 및 정제
        """
        print("📊 데이터 로딩 시작...")
        
        # CSV 파일 로드
        self.df = pd.read_csv(self.csv_file_path)
        print(f"원본 데이터: {len(self.df)}행 {len(self.df.columns)}열")
        
        # 데이터 복사본 생성
        self.cleaned_df = self.df.copy()
        
        # 데이터 정제 작업
        print("\n🧹 데이터 정제 중...")
        
        # 1. 날짜 정제 (과학적 표기법 수정)
        invalid_dates = self.cleaned_df['Date'].str.contains('E\\+', na=False)
        if invalid_dates.any():
            print(f"   - 잘못된 날짜 형식 {invalid_dates.sum()}개 발견 및 수정")
            # 2.00E+05를 2025-09-22로 가정하여 수정
            self.cleaned_df.loc[invalid_dates, 'Date'] = '2025-09-22'
        
        # 2. 날짜 컬럼을 datetime으로 변환
        self.cleaned_df['Date'] = pd.to_datetime(self.cleaned_df['Date'])
        
        # 3. 카테고리와 제품명 정규화 (대소문자 통일)
        self.cleaned_df['Category'] = self.cleaned_df['Category'].str.title()
        self.cleaned_df['ProductName'] = self.cleaned_df['ProductName'].str.title()
        self.cleaned_df['Salesperson'] = self.cleaned_df['Salesperson'].str.title()
        
        # 4. 무효한 제품 데이터 제거
        invalid_products = self.cleaned_df['ProductID'] == 'P0000'
        if invalid_products.any():
            print(f"   - 무효한 제품 데이터 {invalid_products.sum()}개 제거")
            self.cleaned_df = self.cleaned_df[~invalid_products]
        
        # 5. 빈 Quantity 값 처리
        missing_qty = self.cleaned_df['Quantity'].isna()
        if missing_qty.any():
            print(f"   - 빈 수량 값 {missing_qty.sum()}개 발견")
            # 해당 제품의 평균 수량으로 대체
            for idx in self.cleaned_df[missing_qty].index:
                product_id = self.cleaned_df.loc[idx, 'ProductID']
                avg_qty = self.cleaned_df[
                    (self.cleaned_df['ProductID'] == product_id) & 
                    (self.cleaned_df['Quantity'].notna())
                ]['Quantity'].mean()
                if not pd.isna(avg_qty):
                    self.cleaned_df.loc[idx, 'Quantity'] = round(avg_qty)
                    print(f"     - {product_id}: 평균 수량 {round(avg_qty)}로 대체")
        
        # 6. TotalPrice 재계산
        self.cleaned_df['TotalPrice'] = self.cleaned_df['Quantity'] * self.cleaned_df['UnitPrice']
        
        # 7. 빈 Salesperson 값 처리
        missing_sales = self.cleaned_df['Salesperson'].isna() | (self.cleaned_df['Salesperson'] == '')
        if missing_sales.any():
            print(f"   - 빈 영업사원 정보 {missing_sales.sum()}개를 'Unknown'으로 처리")
            self.cleaned_df.loc[missing_sales, 'Salesperson'] = 'Unknown'
        
        print(f"정제된 데이터: {len(self.cleaned_df)}행")
        print("✅ 데이터 정제 완료\n")
        
    def analyze_data(self):
        """
        데이터 분석 수행
        """
        print("📈 데이터 분석 시작...")
        
        # cleaned_df가 None이 아님을 확인
        if self.cleaned_df is None:
            raise ValueError("cleaned_df가 초기화되지 않았습니다. load_and_clean_data()를 먼저 실행하세요.")
        
        # 기본 통계
        self.analysis_results['basic_stats'] = {
            'total_sales': self.cleaned_df['TotalPrice'].sum(),
            'total_quantity': self.cleaned_df['Quantity'].sum(),
            'avg_order_value': self.cleaned_df['TotalPrice'].mean(),
            'total_orders': len(self.cleaned_df),
            'date_range': {
                'start': self.cleaned_df['Date'].min(),
                'end': self.cleaned_df['Date'].max()
            }
        }
        
        # 카테고리별 분석
        category_analysis = self.cleaned_df.groupby('Category').agg({
            'TotalPrice': ['sum', 'mean', 'count'],
            'Quantity': 'sum'
        }).round(2)
        category_analysis.columns = ['총매출', '평균주문금액', '주문수', '총판매량']
        self.analysis_results['category_analysis'] = category_analysis.sort_values('총매출', ascending=False)
        
        # 지역별 분석
        region_analysis = self.cleaned_df.groupby('Region').agg({
            'TotalPrice': ['sum', 'mean', 'count'],
            'Quantity': 'sum'
        }).round(2)
        region_analysis.columns = ['총매출', '평균주문금액', '주문수', '총판매량']
        self.analysis_results['region_analysis'] = region_analysis.sort_values('총매출', ascending=False)
        
        # 영업사원별 분석
        salesperson_analysis = self.cleaned_df.groupby('Salesperson').agg({
            'TotalPrice': ['sum', 'mean', 'count'],
            'Quantity': 'sum'
        }).round(2)
        salesperson_analysis.columns = ['총매출', '평균주문금액', '주문수', '총판매량']
        self.analysis_results['salesperson_analysis'] = salesperson_analysis.sort_values('총매출', ascending=False)
        
        # 제품별 분석 (상위 10개)
        product_analysis = self.cleaned_df.groupby(['ProductID', 'ProductName']).agg({
            'TotalPrice': ['sum', 'mean', 'count'],
            'Quantity': 'sum'
        }).round(2)
        product_analysis.columns = ['총매출', '평균주문금액', '주문수', '총판매량']
        self.analysis_results['top_products'] = product_analysis.sort_values('총매출', ascending=False).head(10)
        
        # 일별 매출 트렌드
        daily_sales = self.cleaned_df.groupby('Date').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).round(2)
        self.analysis_results['daily_trends'] = daily_sales
        
        # 주간별 분석 (주차별)
        self.cleaned_df['Week'] = self.cleaned_df['Date'].dt.isocalendar().week
        weekly_analysis = self.cleaned_df.groupby('Week').agg({
            'TotalPrice': 'sum',
            'Quantity': 'sum'
        }).round(2)
        self.analysis_results['weekly_analysis'] = weekly_analysis
        self.analysis_results['weekly_analysis'] = weekly_analysis
        
        print("✅ 데이터 분석 완료\n")
        
    def generate_word_report(self) -> str:
        """
        워드 문서 형태의 보고서 생성
        """
        print("📄 워드 보고서 생성 중...")
        
        # analysis_results가 있는지 확인
        if not self.analysis_results:
            raise ValueError("분석 결과가 없습니다. analyze_data()를 먼저 실행하세요.")
        
        # 새 워드 문서 생성
        doc = Document()
        
        # 문서 제목
        title = doc.add_heading('제품 판매 데이터 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 기본 통계 가져오기
        basic_stats = self.analysis_results['basic_stats']
        
        # 보고서 기본 정보
        doc.add_paragraph()
        info_para = doc.add_paragraph()
        info_para.add_run(f"분석 기간: ").bold = True
        info_para.add_run(f"{basic_stats['date_range']['start'].strftime('%Y년 %m월 %d일')} ~ {basic_stats['date_range']['end'].strftime('%Y년 %m월 %d일')}")
        
        info_para2 = doc.add_paragraph()
        info_para2.add_run(f"보고서 생성일: ").bold = True
        info_para2.add_run(f"{datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
        
        # 전체 요약 섹션
        doc.add_heading('📊 전체 요약', level=1)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        summary_data = [
            ('📊 총 매출', f"${basic_stats['total_sales']:,.2f}"),
            ('📦 총 판매량', f"{basic_stats['total_quantity']:,.0f}개"),
            ('💰 평균 주문금액', f"${basic_stats['avg_order_value']:,.2f}"),
            ('🛒 총 주문수', f"{basic_stats['total_orders']:,}건"),
            ('📅 분석 기간', f"{(basic_stats['date_range']['end'] - basic_stats['date_range']['start']).days + 1}일")
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
            # 헤더 셀 굵게
            summary_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # 주요 인사이트 섹션
        doc.add_heading('💡 주요 인사이트', level=1)
        insights = [
            f"매출 리더: {self.analysis_results['category_analysis'].index[0]} 카테고리가 ${self.analysis_results['category_analysis']['총매출'].iloc[0]:,.2f}로 최고 매출 기록",
            f"지역 성과: {self.analysis_results['region_analysis'].index[0]} 지역이 ${self.analysis_results['region_analysis']['총매출'].iloc[0]:,.2f}로 최고 실적",
            f"영업 성과: {self.analysis_results['salesperson_analysis'].index[0]} 영업사원이 ${self.analysis_results['salesperson_analysis']['총매출'].iloc[0]:,.2f} 매출로 1위",
            f"인기 제품: {self.analysis_results['top_products'].index[0][1]}이 ${self.analysis_results['top_products']['총매출'].iloc[0]:,.2f} 매출로 베스트셀러"
        ]
        
        for insight in insights:
            p = doc.add_paragraph()
            p.add_run("• ").bold = True
            p.add_run(insight)
        
        # 카테고리별 분석 섹션
        doc.add_heading('📈 카테고리별 분석', level=1)
        
        category_table = doc.add_table(rows=len(self.analysis_results['category_analysis']) + 1, cols=5)
        category_table.style = 'Table Grid'
        category_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 테이블 헤더
        hdr_cells = category_table.rows[0].cells
        headers = ['카테고리', '총 매출', '평균 주문금액', '주문 수', '총 판매량']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # 데이터 행 추가
        for i, (category, row) in enumerate(self.analysis_results['category_analysis'].iterrows()):
            cells = category_table.rows[i + 1].cells
            cells[0].text = category
            cells[1].text = f"${row['총매출']:,.2f}"
            cells[2].text = f"${row['평균주문금액']:,.2f}"
            cells[3].text = f"{row['주문수']:,.0f}"
            cells[4].text = f"{row['총판매량']:,.0f}"
        
        # 지역별 분석 섹션
        doc.add_heading('🌍 지역별 분석', level=1)
        
        region_table = doc.add_table(rows=len(self.analysis_results['region_analysis']) + 1, cols=5)
        region_table.style = 'Table Grid'
        region_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 테이블 헤더
        hdr_cells = region_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header.replace('카테고리', '지역')
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # 데이터 행 추가
        for i, (region, row) in enumerate(self.analysis_results['region_analysis'].iterrows()):
            cells = region_table.rows[i + 1].cells
            cells[0].text = region
            cells[1].text = f"${row['총매출']:,.2f}"
            cells[2].text = f"${row['평균주문금액']:,.2f}"
            cells[3].text = f"{row['주문수']:,.0f}"
            cells[4].text = f"{row['총판매량']:,.0f}"
        
        # 영업사원별 성과 섹션
        doc.add_heading('👨‍💼 영업사원별 성과', level=1)
        
        sales_table = doc.add_table(rows=len(self.analysis_results['salesperson_analysis']) + 1, cols=5)
        sales_table.style = 'Table Grid'
        sales_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 테이블 헤더
        hdr_cells = sales_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header.replace('카테고리', '영업사원')
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # 데이터 행 추가
        for i, (salesperson, row) in enumerate(self.analysis_results['salesperson_analysis'].iterrows()):
            cells = sales_table.rows[i + 1].cells
            cells[0].text = salesperson
            cells[1].text = f"${row['총매출']:,.2f}"
            cells[2].text = f"${row['평균주문금액']:,.2f}"
            cells[3].text = f"{row['주문수']:,.0f}"
            cells[4].text = f"{row['총판매량']:,.0f}"
        
        # 상위 제품 분석 섹션 (Top 10)
        doc.add_heading('🏆 상위 제품 (Top 10)', level=1)
        
        product_table = doc.add_table(rows=len(self.analysis_results['top_products']) + 1, cols=6)
        product_table.style = 'Table Grid'
        product_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 테이블 헤더
        hdr_cells = product_table.rows[0].cells
        product_headers = ['제품 ID', '제품명', '총 매출', '평균 주문금액', '주문 수', '총 판매량']
        for i, header in enumerate(product_headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # 데이터 행 추가
        for i, ((product_id, product_name), row) in enumerate(self.analysis_results['top_products'].iterrows()):
            cells = product_table.rows[i + 1].cells
            cells[0].text = product_id
            cells[1].text = product_name
            cells[2].text = f"${row['총매출']:,.2f}"
            cells[3].text = f"${row['평균주문금액']:,.2f}"
            cells[4].text = f"{row['주문수']:,.0f}"
            cells[5].text = f"{row['총판매량']:,.0f}"
        
        # 데이터 품질 이슈 섹션
        doc.add_heading('⚠️ 데이터 품질 이슈', level=1)
        
        quality_issues = [
            "일부 제품의 수량 정보가 누락되어 해당 제품의 평균값으로 대체했습니다.",
            "일부 영업사원 정보가 누락되어 'Unknown'으로 처리했습니다.",
            "잘못된 날짜 형식(과학적 표기법) 1건을 수정했습니다.",
            "무효한 제품 데이터(P0000) 1건을 제거했습니다."
        ]
        
        for issue in quality_issues:
            p = doc.add_paragraph()
            p.add_run("• ").bold = True
            p.add_run(issue)
        
        # 문서 하단
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run("보고서 생성 시간: ").bold = True
        footer_para.add_run(datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분'))
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 워드 파일 저장
        docx_file_path = self.csv_file_path.replace('.csv', '_sales_report.docx')
        doc.save(docx_file_path)
        
        print(f"✅ 워드 보고서 저장: {docx_file_path}")
        print("✅ 워드 보고서 생성 완료\n")
        
        return docx_file_path
        
    def send_email_with_report(self, 
                             docx_file_path: str,
                             sender_email: str,
                             sender_password: str,
                             recipient_emails: list,
                             smtp_server: str = "smtp.gmail.com",
                             smtp_port: int = 587,
                             subject: Optional[str] = None) -> bool:
        """
        워드 보고서를 첨부하여 이메일 전송
        
        Args:
            docx_file_path: 워드 파일 경로
            sender_email: 발신자 이메일
            sender_password: 발신자 이메일 비밀번호 (앱 비밀번호 권장)
            recipient_emails: 수신자 이메일 리스트
            smtp_server: SMTP 서버 주소
            smtp_port: SMTP 포트
            subject: 이메일 제목 (기본값: 자동 생성)
            
        Returns:
            bool: 전송 성공 여부
        """
        print("📧 이메일 전송 준비 중...")
        
        try:
            # 파일 존재 확인
            if not os.path.exists(docx_file_path):
                raise FileNotFoundError(f"워드 파일을 찾을 수 없습니다: {docx_file_path}")
            
            # 기본 제목 설정
            if subject is None:
                basic_stats = self.analysis_results.get('basic_stats', {})
                if basic_stats and 'date_range' in basic_stats:
                    start_date = basic_stats['date_range']['start'].strftime('%Y-%m-%d')
                    end_date = basic_stats['date_range']['end'].strftime('%Y-%m-%d')
                    subject = f"📊 제품 판매 데이터 분석 보고서 ({start_date} ~ {end_date})"
                else:
                    subject = f"📊 제품 판매 데이터 분석 보고서 - {datetime.now().strftime('%Y-%m-%d')}"
            
            # 이메일 본문 작성
            email_body = self._create_email_body()
            
            # 멀티파트 메시지 생성
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = ', '.join(recipient_emails)
            msg['Subject'] = subject
            
            # 본문 추가
            msg.attach(MIMEText(email_body, 'html', 'utf-8'))
            
            # 워드 파일 첨부
            with open(docx_file_path, 'rb') as attachment:
                part = MIMEApplication(attachment.read(), _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document')
                filename = os.path.basename(docx_file_path)
                part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
                msg.attach(part)
            
            # SMTP 서버 연결 및 전송
            print(f"📤 SMTP 서버 연결 중... ({smtp_server}:{smtp_port})")
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()  # TLS 암호화 시작
            
            print("🔐 로그인 중...")
            server.login(sender_email, sender_password)
            
            print(f"📨 이메일 전송 중... (수신자: {len(recipient_emails)}명)")
            text = msg.as_string()
            server.sendmail(sender_email, recipient_emails, text)
            server.quit()
            
            print("✅ 이메일 전송 완료!")
            print(f"   - 수신자: {', '.join(recipient_emails)}")
            print(f"   - 제목: {subject}")
            print(f"   - 첨부파일: {filename}")
            
            return True
            
        except Exception as e:
            print(f"❌ 이메일 전송 실패: {str(e)}")
            print("💡 확인사항:")
            print("   - Gmail의 경우 앱 비밀번호를 사용하세요")
            print("   - 2단계 인증이 활성화되어 있는지 확인하세요")
            print("   - SMTP 설정이 올바른지 확인하세요")
            return False
    
    def _create_email_body(self) -> str:
        """
        이메일 본문 HTML 생성
        """
        basic_stats = self.analysis_results.get('basic_stats', {})
        
        if not basic_stats:
            return "<p>분석 보고서가 첨부되어 있습니다.</p>"
        
        # 날짜 정보
        if 'date_range' in basic_stats:
            start_date = basic_stats['date_range']['start'].strftime('%Y년 %m월 %d일')
            end_date = basic_stats['date_range']['end'].strftime('%Y년 %m월 %d일')
            date_info = f"{start_date} ~ {end_date}"
        else:
            date_info = "분석 기간 정보 없음"
        
        # 주요 인사이트
        insights_html = ""
        if 'category_analysis' in self.analysis_results:
            top_category = self.analysis_results['category_analysis'].index[0]
            top_category_sales = self.analysis_results['category_analysis']['총매출'].iloc[0]
            insights_html += f"<li>🏆 최고 매출 카테고리: <strong>{top_category}</strong> (${top_category_sales:,.2f})</li>"
        
        if 'region_analysis' in self.analysis_results:
            top_region = self.analysis_results['region_analysis'].index[0]
            top_region_sales = self.analysis_results['region_analysis']['총매출'].iloc[0]
            insights_html += f"<li>🌟 최고 성과 지역: <strong>{top_region}</strong> (${top_region_sales:,.2f})</li>"
        
        if 'salesperson_analysis' in self.analysis_results:
            top_salesperson = self.analysis_results['salesperson_analysis'].index[0]
            top_sales_amount = self.analysis_results['salesperson_analysis']['총매출'].iloc[0]
            insights_html += f"<li>👑 최고 성과 영업사원: <strong>{top_salesperson}</strong> (${top_sales_amount:,.2f})</li>"
        
        html_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px;">
                    📊 제품 판매 데이터 분석 보고서
                </h2>
                
                <p>안녕하세요,</p>
                <p><strong>{date_info}</strong> 기간의 제품 판매 데이터 분석 보고서를 전송드립니다.</p>
                
                <div style="background-color: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0;">
                    <h3 style="color: #2c3e50; margin-top: 0;">📈 주요 분석 결과</h3>
                    <ul style="padding-left: 20px;">
                        <li>💰 총 매출: <strong>${basic_stats.get('total_sales', 0):,.2f}</strong></li>
                        <li>📦 총 판매량: <strong>{basic_stats.get('total_quantity', 0):,.0f}개</strong></li>
                        <li>🛒 총 주문수: <strong>{basic_stats.get('total_orders', 0):,}건</strong></li>
                        <li>💵 평균 주문금액: <strong>${basic_stats.get('avg_order_value', 0):,.2f}</strong></li>
                    </ul>
                </div>
                
                <div style="background-color: #e8f4fd; padding: 20px; border-radius: 8px; margin: 20px 0;">
                    <h3 style="color: #2c3e50; margin-top: 0;">💡 핵심 인사이트</h3>
                    <ul style="padding-left: 20px;">
                        {insights_html}
                    </ul>
                </div>
                
                <div style="background-color: #fff3cd; padding: 15px; border-radius: 8px; margin: 20px 0; border-left: 4px solid #ffc107;">
                    <p style="margin: 0;"><strong>📎 첨부파일:</strong> 상세한 분석 내용이 포함된 워드 문서가 첨부되어 있습니다.</p>
                </div>
                
                <hr style="border: none; border-top: 1px solid #ddd; margin: 30px 0;">
                
                <p style="color: #666; font-size: 14px;">
                    이 보고서는 자동으로 생성되었습니다.<br>
                    생성 시간: {datetime.now().strftime('%Y년 %m월 %d일 %H시 %M분')}<br>
                    문의사항이 있으시면 데이터 분석팀으로 연락해주세요.
                </p>
            </div>
        </body>
        </html>
        """
        
        return html_body
        
    def run_full_analysis(self) -> Dict[str, Any]:
        """
        전체 분석 프로세스 실행 (워드 보고서만 생성)
        """
        print("🚀 판매 데이터 분석 시작!\n")
        
        self.load_and_clean_data()
        self.analyze_data()
        docx_path = self.generate_word_report()
        
        print("🎉 분석 완료!")
        print(f"📄 워드 보고서: {docx_path}")
        
        return {
            'word_report': docx_path,
            'analysis_results': self.analysis_results
        }


def main() -> None:
    """
    메인 실행 함수 - 워드 보고서 생성 및 이메일 전송 옵션
    """
    csv_file = "references/cicd_data.csv"
    
    # 분석기 생성 및 실행
    analyzer = SalesDataAnalyzer(csv_file)
    results = analyzer.run_full_analysis()
    
    print("\n" + "="*50)
    print("분석 결과 파일:")
    print(f"  워드 보고서: {results['word_report']}")
    
    # 환경변수에서 이메일 설정 확인 (GitHub Actions용)
    sender_email = os.getenv('SENDER_EMAIL')
    email_password = os.getenv('EMAIL_PASSWORD')
    recipient_email = os.getenv('RECIPIENT_EMAIL')
    
    if sender_email and email_password and recipient_email:
        print("\n📧 환경변수에서 이메일 설정을 찾았습니다. 자동으로 이메일을 전송합니다...")
        print(f"🔍 디버깅: 발신자 이메일 - {sender_email[:3]}***@{sender_email.split('@')[1] if '@' in sender_email else 'unknown'}")
        print(f"🔍 디버깅: 수신자 이메일 - {recipient_email[:3]}***")
        
        # 수신자 이메일 처리 (쉼표로 구분된 경우)
        recipient_emails = [email.strip() for email in recipient_email.split(',') if email.strip()]
        
        # 이메일 주소 유효성 검사
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        
        if not re.match(email_pattern, sender_email):
            print(f"❌ 발신자 이메일 형식이 올바르지 않습니다: {sender_email}")
            return
            
        for email in recipient_emails:
            if not re.match(email_pattern, email):
                print(f"❌ 수신자 이메일 형식이 올바르지 않습니다: {email}")
                return
        
        # 이메일 전송
        success = analyzer.send_email_with_report(
            docx_file_path=results['word_report'],
            sender_email=sender_email,
            sender_password=email_password,
            recipient_emails=recipient_emails,
            subject=f"📊 주간 판매 데이터 분석 보고서 - {datetime.now().strftime('%Y-%m-%d')}"
        )
        
        if success:
            print("🎉 자동 이메일 전송이 완료되었습니다!")
        else:
            print("⚠️ 보고서는 생성되었지만 이메일 전송에 실패했습니다.")
        return
    
    # 환경변수가 없는 경우 대화형 모드
    print("\n" + "="*50)
    send_email = input("📧 이메일로 보고서를 전송하시겠습니까? (y/n): ").lower().strip()
    
    if send_email in ['y', 'yes', '예', 'ㅇ']:
        print("\n📧 이메일 설정을 입력해주세요:")
        
        # 이메일 설정 입력
        sender_email = input("발신자 이메일: ")
        sender_password = input("발신자 비밀번호 (Gmail 앱 비밀번호 권장): ")
        
        recipient_input = input("수신자 이메일 (여러 명은 쉼표로 구분): ")
        recipient_emails = [email.strip() for email in recipient_input.split(',') if email.strip()]
        
        # 선택적 설정
        custom_subject = input("이메일 제목 (엔터 시 기본값 사용): ").strip()
        subject = custom_subject if custom_subject else None
        
        smtp_input = input("SMTP 서버 (엔터 시 Gmail 기본값): ").strip()
        smtp_server = smtp_input if smtp_input else "smtp.gmail.com"
        
        port_input = input("SMTP 포트 (엔터 시 587): ").strip()
        smtp_port = int(port_input) if port_input.isdigit() else 587
        
        print("\n📤 이메일 전송 중...")
        
        # 이메일 전송
        success = analyzer.send_email_with_report(
            docx_file_path=results['word_report'],
            sender_email=sender_email,
            sender_password=sender_password,
            recipient_emails=recipient_emails,
            smtp_server=smtp_server,
            smtp_port=smtp_port,
            subject=subject
        )
        
        if success:
            print("🎉 모든 작업이 완료되었습니다!")
        else:
            print("⚠️ 보고서는 생성되었지만 이메일 전송에 실패했습니다.")
    else:
        print("📄 보고서 생성이 완료되었습니다.")


def send_report_email_example():
    """
    이메일 전송 예시 함수 (참고용)
    """
    csv_file = "references/cicd_data.csv"
    
    # 분석 실행
    analyzer = SalesDataAnalyzer(csv_file)
    results = analyzer.run_full_analysis()
    
    # 이메일 전송 예시
    success = analyzer.send_email_with_report(
        docx_file_path=results['word_report'],
        sender_email="ilhj1228@gmail.com",  # 여기에 실제 이메일 입력
        sender_password="clqq xbqj jbzg nzjy",   # 여기에 앱 비밀번호 입력
        recipient_emails=[""],
        subject="📊 월간 판매 분석 보고서"
    )
    
    return success


if __name__ == "__main__":
    main()